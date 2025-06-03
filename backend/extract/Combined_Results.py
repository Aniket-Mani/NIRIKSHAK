

"""
Processes all students within a specific professoruploads document,
generates individual marksheets for each, saves them to the Results collection,
and then generates a combined class marksheet.
Updates the professoruploads document with the status and combined report ID.
"""
import os
# THIS MUST BE AT THE VERY TOP, BEFORE ANY OTHER IMPORTS
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import json
import re
import sys
import traceback  # For printing full tracebacks
from typing import Dict, List, Union, Any, Tuple
import pandas as pd
import numpy as np
from pymongo import MongoClient, UpdateOne, errors as PyMongoErrors
from bson.objectid import ObjectId
import gridfs
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Corrected import
import subprocess
import argparse
from datetime import datetime, timezone
# from dotenv import load_dotenv # Uncomment if you use a .env file locally

# ---------------------------------------------------------------------------
# CONFIGURATION
# ---------------------------------------------------------------------------
MONGO_CONNECTION_STRING = os.getenv("MONGO_CONNECTION_STRING", "mongodb://localhost:27017/")
DATABASE_NAME = os.getenv("MONGO_DB_NAME", "smart")

LOGO_IMAGE_FILENAME = os.getenv("LOGO_IMAGE_FILENAME", "logo_name.png") # Ensure this is the correct filename
try:
    PROJECT_ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
except NameError:
    PROJECT_ROOT_DIR = os.path.abspath(os.path.join(os.getcwd(), '../..'))

LOGO_IMAGE_PATH = os.path.join(PROJECT_ROOT_DIR, "backend", "logo", LOGO_IMAGE_FILENAME)

OUTPUT_DIR_COMBINED = os.path.join(PROJECT_ROOT_DIR, "backend", "uploads", "generated_combined_marksheets_temp")
os.makedirs(OUTPUT_DIR_COMBINED, exist_ok=True)

INDIVIDUAL_RESULTS_BUCKET_NAME = "results_marksheets"
CLASS_AGGREGATE_BUCKET_NAME = "class_aggregate_reports"

PROFESSOR_UPLOADS_COLLECTION_NAME = "professoruploads"
RESULTS_COLLECTION_NAME = "Results"

# ---------------------------------------------------------------------------
# DATABASE CONNECTIONS
# ---------------------------------------------------------------------------
try:
    print(f"INFO (CombinedResults): Connecting to MongoDB at {MONGO_CONNECTION_STRING}...", file=sys.stderr)
    client = MongoClient(MONGO_CONNECTION_STRING, serverSelectionTimeoutMS=5000)
    client.admin.command('ping')
    db = client[DATABASE_NAME]
    professoruploads_collection = db[PROFESSOR_UPLOADS_COLLECTION_NAME]
    results_collection = db[RESULTS_COLLECTION_NAME]
    fs_individual_results_bucket = gridfs.GridFS(db, collection=INDIVIDUAL_RESULTS_BUCKET_NAME)
    fs_class_aggregate_bucket = gridfs.GridFS(db, collection=CLASS_AGGREGATE_BUCKET_NAME)
    print(f"INFO (CombinedResults): Connected to MongoDB: db='{DATABASE_NAME}'", file=sys.stderr)
except PyMongoErrors.ServerSelectionTimeoutError as e: # More specific error
    print(f"FATAL (CombinedResults): Could not connect to MongoDB (Timeout). Error: {e}", file=sys.stderr)
    sys.exit(1)
except PyMongoErrors.ConnectionFailure as e:
    print(f"FATAL (CombinedResults): Could not connect to MongoDB. Error: {e}", file=sys.stderr)
    sys.exit(1)
except Exception as e:
    print(f"FATAL (CombinedResults): An unexpected error occurred during MongoDB setup. Error: {e}", file=sys.stderr)
    traceback.print_exc(file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# EMBEDDING MODEL & HELPERS
# ---------------------------------------------------------------------------
try:
    print("INFO (CombinedResults): Loading sentence embedding model 'all-MiniLM-L6-v2'...", file=sys.stderr)
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    print("INFO (CombinedResults): Sentence embedding model loaded.", file=sys.stderr)
except Exception as e:
    print(f"FATAL (CombinedResults): Could not load SentenceTransformer model. Error: {e}", file=sys.stderr)
    traceback.print_exc(file=sys.stderr)
    sys.exit(1)


def preprocess(text: str) -> str:
    if not text or not isinstance(text, str):
        return ""
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def embed(text_to_embed: str) -> Union[np.ndarray, None]:
    processed_text = preprocess(text_to_embed)
    return embedding_model.encode(processed_text, convert_to_numpy=True) if processed_text else None


def cos_sim(v1: Union[np.ndarray, None], v2: Union[np.ndarray, None]) -> float:
    if v1 is None or v2 is None or v1.size == 0 or v2.size == 0:
        return 0.0
    v1_r = v1.reshape(1, -1) if v1.ndim == 1 else v1
    v2_r = v2.reshape(1, -1) if v2.ndim == 1 else v2
    if v1_r.shape[1] != v2_r.shape[1]:
        print(
            f"WARNING (CombinedResults): Cosine similarity dim mismatch. v1: {v1_r.shape}, v2: {v2_r.shape}",
            file=sys.stderr
        )
        return 0.0
    return float(cosine_similarity(v1_r, v2_r)[0][0])

# ---------------------------------------------------------------------------
# DATA PARSING
# ---------------------------------------------------------------------------
def parse_reference_answers_from_processed_json(professor_processed_json_list: List[Dict[str, Any]]) -> Dict[str, Any]:
    items = []
    for q_idx, q_data in enumerate(professor_processed_json_list):
        if not isinstance(q_data, dict):
            print(f"WARNING (CombinedResults): Item in processedJSON is not a dict: {q_data}", file=sys.stderr)
            continue
        question_id_from_prof = q_data.get('questionNo', f"q_default_{q_idx + 1}")
        temp_id = str(question_id_from_prof).strip()
        match = re.match(r"^[Qq]([0-9]+[a-zA-Z]?)$", temp_id)
        if match:
            qid_normalized = match.group(1)
        else:
            qid_normalized = temp_id.replace(" ", "").replace(".", "_").replace("/", "_")

        ref_answers_list = q_data.get("Answers", [])
        actual_refs = [str(ans) if ans is not None else "" for ans in ref_answers_list]
        padded_refs = (actual_refs + ["", "", ""])[:3]

        items.append({
            "question_id": qid_normalized,
            "max_marks": int(q_data.get("marks", 0)),
            "answer1": padded_refs[0],
            "answer2": padded_refs[1],
            "answer3": padded_refs[2],
            "question_text": str(q_data.get("questionText", ""))
        })
    return {"questions": items}


def build_reference_vectors(parsed_ref_answers: Dict[str, Any]) -> Dict[str, List[Union[np.ndarray, None]]]:
    vector_cache = {}
    for q in parsed_ref_answers.get("questions", []):
        vector_cache[q["question_id"]] = [
            embed(q["answer1"]),
            embed(q["answer2"]),
            embed(q["answer3"])
        ]
    return vector_cache

# ---------------------------------------------------------------------------
# SIMILARITY CALCULATION
# ---------------------------------------------------------------------------
def calculate_similarity_for_student(
    student_data_from_prof_doc: Dict[str, Any],
    ref_vecs: Dict[str, List[Union[np.ndarray, None]]],
    max_marks_map: Dict[str, int]
) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    roll_no = student_data_from_prof_doc.get("roll_no")
    if not roll_no:
        print("WARNING (CombinedResults): Student entry in professoruploads missing 'roll_no'. Skipping.", file=sys.stderr)
        return pd.DataFrame()

    student_answers_list = student_data_from_prof_doc.get("answers", [])
    if not isinstance(student_answers_list, list):
        print(
            f"WARNING (CombinedResults): 'answers' for roll {roll_no} (from professoruploads) is not a list or missing. Skipping.",
            file=sys.stderr
        )
        return pd.DataFrame()

    for ans_data in student_answers_list:
        if not isinstance(ans_data, dict):
            continue
        qid_from_student = ans_data.get("question_no")
        if not qid_from_student:
            continue

        temp_id_student = str(qid_from_student).strip()
        match_student = re.match(r"^[Qq]?([0-9]+[a-zA-Z]?)$", temp_id_student)
        if match_student:
            qid_normalized_student = match_student.group(1)
        else:
            qid_normalized_student = temp_id_student.replace(" ", "").replace(".", "_").replace("/", "_")

        stu_answer_text = ans_data.get("answer_text")
        if stu_answer_text is None:
            continue

        stu_vec = embed(str(stu_answer_text))
        similarity = 0.0
        if stu_vec is not None:
            ref_q_vecs_for_qid = ref_vecs.get(qid_normalized_student, [])
            if not ref_q_vecs_for_qid and qid_normalized_student in max_marks_map:
                print(f"WARNING (CombinedResults): No ref vecs for QID '{qid_normalized_student}' for roll {roll_no}.", file=sys.stderr)

            sim_scores = [cos_sim(stu_vec, ref_v) for ref_v in ref_q_vecs_for_qid if ref_v is not None]
            similarity = max(sim_scores) if sim_scores else 0.0

        question_max_marks = max_marks_map.get(qid_normalized_student, 0)
        if qid_normalized_student not in max_marks_map:
            print(
                f"INFO (CombinedResults): Max marks not found for student QID '{qid_normalized_student}' "
                f"(orig: '{qid_from_student}') for roll {roll_no}. "
                f"Avail ref QIDs: {list(max_marks_map.keys())}",
                file=sys.stderr
            )

        summary_text = str(stu_answer_text)
        student_answer_summary = (summary_text[:100] + "...") if len(summary_text) > 100 else summary_text

        rows.append({
            "roll_no": roll_no,
            "question_id": qid_normalized_student,
            "max_marks": question_max_marks,
            "similarity": round(similarity, 3),
            "student_answer_summary": student_answer_summary
        })

    df = pd.DataFrame(rows)
    if df.empty:
        print(f"INFO (CombinedResults): Similarity dataframe empty for roll {roll_no} (from professoruploads).", file=sys.stderr)
        return pd.DataFrame(columns=["roll_no", "question_id", "max_marks", "similarity", "student_answer_summary", "score"])

    def score_rule(sim: float, max_m: int) -> int:
        sim = float(sim) if sim is not None else 0.0
        max_m = int(max_m) if max_m is not None else 0
        if sim >= 0.90: pct = 1.00
        elif sim >= 0.80: pct = 0.90
        elif sim >= 0.70: pct = 0.75
        elif sim >= 0.60: pct = 0.65
        elif sim >= 0.50: pct = 0.60
        elif sim >= 0.45: pct = 0.50
        else: pct = 0.0
        return int(round(pct * max_m))

    df["score"] = df.apply(lambda r: score_rule(r["similarity"], r["max_marks"]), axis=1)
    return df

# ---------------------------------------------------------------------------
# PDF BUILDERS
# ---------------------------------------------------------------------------
def convert_docx_to_pdf_pandoc(docx_path: str, pdf_path: str) -> bool:
    try:
        pandoc_command = ['pandoc', docx_path, '-o', pdf_path]
        process = subprocess.run(pandoc_command, capture_output=True, text=True, check=False)

        if process.returncode == 0:
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"INFO (CombinedResults): PDF '{pdf_path}' created successfully via Pandoc.", file=sys.stderr)
                return True
            else:
                print(f"ERROR (CombinedResults): Pandoc for '{docx_path}' succeeded (code 0) but PDF '{pdf_path}' not created or is empty.", file=sys.stderr)
                print(f"Pandoc stdout:\n{process.stdout}", file=sys.stderr)
                print(f"Pandoc stderr:\n{process.stderr}", file=sys.stderr)
                return False
        else:
            print(f"ERROR (CombinedResults): Pandoc conversion failed for '{docx_path}'. Code: {process.returncode}", file=sys.stderr)
            print(f"Pandoc stdout:\n{process.stdout}", file=sys.stderr)
            print(f"Pandoc stderr:\n{process.stderr}", file=sys.stderr)
            return False
    except FileNotFoundError:
        print("ERROR (CombinedResults): Pandoc not found. Ensure it's installed and in PATH.", file=sys.stderr)
        return False
    except Exception as e:
        print(f"ERROR (CombinedResults): Pandoc conversion exception for '{docx_path}'. {e}", file=sys.stderr)
        return False

def build_student_pdf(
    df_student_scores: pd.DataFrame,
    roll_no: str,
    logo_image_path_param: str,
    exam_details_for_pdf: Dict[str, Any]
) -> Tuple[Union[str, None], str]:
    if df_student_scores.empty:
        print(f"WARNING (CombinedResults): No graded answers df for student {roll_no} to build PDF. CSV may be empty.", file=sys.stderr)
        csv_name = os.path.join(OUTPUT_DIR_COMBINED, f"{roll_no}_individual_marksheet.csv")
        pd.DataFrame(columns=["question_id", "max_marks", "score", "percentage"]).to_csv(csv_name, index=False)
        return None, csv_name

    df_student_scores['max_marks'] = pd.to_numeric(df_student_scores['max_marks'], errors='coerce').fillna(0).astype(int)
    df_student_scores['score'] = pd.to_numeric(df_student_scores['score'], errors='coerce').fillna(0).astype(int)
    df_student_scores["percentage"] = df_student_scores.apply(
        lambda r: round((r["score"] / r["max_marks"]) * 100, 2) if r["max_marks"] > 0 else 0.0, axis=1
    )

    total_score = df_student_scores["score"].sum()
    total_max_for_percentage = exam_details_for_pdf.get("total_max_marks_from_prof", df_student_scores["max_marks"].sum())
    if total_max_for_percentage == 0 and not df_student_scores.empty:
        total_max_for_percentage = df_student_scores["max_marks"].sum()

    total_pct = round((total_score / total_max_for_percentage) * 100, 2) if total_max_for_percentage > 0 else 0.0

    summary_data = [{"question_id": "Total", "max_marks": total_max_for_percentage, "score": total_score, "percentage": total_pct}]
    summary_df = pd.DataFrame(summary_data)

    cols_for_sheet = ["question_id", "max_marks", "score", "percentage"]
    df_display = df_student_scores[cols_for_sheet].astype(str)
    summary_display_df = summary_df[cols_for_sheet].astype(str)
    sheet_for_display = pd.concat([df_display, summary_display_df], ignore_index=True)
    sheet_for_csv = pd.concat([df_student_scores[cols_for_sheet], summary_df[cols_for_sheet]], ignore_index=True)

    exam_type_short = str(exam_details_for_pdf.get("examType", "Exam")).replace(" ", "_")[:20]
    subject_code_short = str(exam_details_for_pdf.get("subjectCode", "Sub"))
    course_short = str(exam_details_for_pdf.get("course", "Course")).replace(" ", "_")[:15] # Program like MCA
    section_short = str(exam_details_for_pdf.get("sectionType", "All"))
    base_filename = f"{roll_no}_{course_short}_{subject_code_short}_{exam_type_short}_{section_short}_individual_marksheet"

    csv_name = os.path.join(OUTPUT_DIR_COMBINED, f"{base_filename}.csv")
    pdf_name = os.path.join(OUTPUT_DIR_COMBINED, f"{base_filename}.pdf")
    tmp_docx = os.path.join(OUTPUT_DIR_COMBINED, f"tmp_{base_filename}.docx")

    try:
        sheet_for_csv.to_csv(csv_name, index=False)
    except Exception as e:
        print(f"ERROR (CR): Could not save individual CSV for {roll_no}. {e}", file=sys.stderr)

    doc = Document()
    sections = doc.sections # Apply margins
    for section_doc in sections:
        section_doc.left_margin = Inches(0.75)
        section_doc.right_margin = Inches(0.75)
        section_doc.top_margin = Inches(0.5)
        section_doc.bottom_margin = Inches(0.5)

    # --- New Header Start (Student PDF) ---
    if os.path.exists(logo_image_path_param):
        try:
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_image_path_param, width=Inches(4.0)) # Adjusted "good size"
        except Exception as e_img:
            print(f"WARN (CR Stud PDF): Logo add error for {roll_no}. {e_img}", file=sys.stderr)
            # Optionally add placeholder text if logo fails
            # doc.add_paragraph("[Logo Loading Error]").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        print(f"WARN (CR Stud PDF): Logo missing for {roll_no} at '{logo_image_path_param}'.", file=sys.stderr)
        # Optionally add placeholder text if logo missing
        # doc.add_paragraph("[Logo Not Found]").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    inst_name_p1 = doc.add_paragraph()
    inst_name_p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_inst_name1 = inst_name_p1.add_run("NATIONAL INSTITUTE OF TECHNOLOGY")
    run_inst_name1.font.name = 'Times New Roman'
    run_inst_name1.font.size = Pt(16)
    run_inst_name1.font.bold = True

    inst_name_p2 = doc.add_paragraph()
    inst_name_p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_inst_name2 = inst_name_p2.add_run("TIRUCHIRAPPALLI")
    run_inst_name2.font.name = 'Times New Roman'
    run_inst_name2.font.size = Pt(14)
    run_inst_name2.font.bold = True
    
    doc.add_paragraph() # Spacer paragraph after header
    # --- New Header End (Student PDF) ---

    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tr_run = p_title.add_run("INDIVIDUAL MARKSHEET")
    tr_run.font.name = 'Times New Roman'; tr_run.font.size = Pt(16); tr_run.font.bold = True
    doc.add_paragraph()

    # Student Details Table
    details_data_stud = [
        ("Roll No:", str(roll_no)),
        ("Program:", str(exam_details_for_pdf.get("course", ""))), # Course is Program
        ("Course Name:", str(exam_details_for_pdf.get("subject", ""))), # Subject is Course Name
        ("Course Code:", str(exam_details_for_pdf.get("subjectCode", ""))),
        ("Exam Type:", str(exam_details_for_pdf.get("examType", ""))),
        ("Section:", str(exam_details_for_pdf.get("sectionType", ""))),
        ("Year:", str(exam_details_for_pdf.get("year", ""))),
        ("Semester:", str(exam_details_for_pdf.get("semester", ""))),
    ]
    dt = doc.add_table(rows=len(details_data_stud), cols=2)
    dt.style = 'Table Grid'
    for i, (label, value) in enumerate(details_data_stud):
        label_cell_p_dt = dt.cell(i, 0).paragraphs[0]
        label_cell_p_dt.text = label # Direct assignment is fine, then format runs
        for run in label_cell_p_dt.runs: 
            run.font.name='Times New Roman'; run.font.size=Pt(10); run.font.bold=True
        
        value_cell_p_dt = dt.cell(i, 1).paragraphs[0]
        value_cell_p_dt.text = value # Direct assignment
        for run in value_cell_p_dt.runs: 
            run.font.name='Times New Roman'; run.font.size=Pt(10); run.font.bold=False
    try: # Column widths for student details
        dt.columns[0].width = Inches(2.0)
        dt.columns[1].width = Inches(4.0)
    except IndexError: # Handle case where table might not have columns if details_data_stud is empty
        pass 
    doc.add_paragraph() 

    qwp = doc.add_paragraph()
    qwr = qwp.add_run("Question-wise Marks:")
    qwr.font.name='Times New Roman'; qwr.font.bold = True; qwr.font.size = Pt(12)

    if not sheet_for_display.empty:
        num_cols = len(sheet_for_display.columns)
        mt = doc.add_table(rows=1, cols=num_cols)
        mt.style = 'Table Grid'
        hdr_cells = mt.rows[0].cells
        for i, column_name in enumerate(sheet_for_display.columns):
            cell_p_mt_hdr = hdr_cells[i].paragraphs[0]; cell_p_mt_hdr.text = ''
            run_mt_hdr = cell_p_mt_hdr.add_run(str(column_name))
            run_mt_hdr.font.name='Times New Roman'; run_mt_hdr.font.bold = True; run_mt_hdr.font.size = Pt(9)
            cell_p_mt_hdr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        for _, data_row_series in sheet_for_display.iterrows():
            row_cells = mt.add_row().cells
            for i, cell_value in enumerate(data_row_series):
                cell_p_mt_data = row_cells[i].paragraphs[0]
                cell_p_mt_data.text = str(cell_value)
                for run in cell_p_mt_data.runs: run.font.name='Times New Roman'; run.font.size = Pt(9)
                cell_p_mt_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        doc.add_paragraph("No scores to display.")

    pdf_generated_successfully = False
    try:
        doc.save(tmp_docx)
        pdf_generated_successfully = convert_docx_to_pdf_pandoc(tmp_docx, pdf_name)
    except Exception as e_conv:
        print(f"ERROR (CR): DOCX to PDF failed for student {roll_no}. {e_conv}", file=sys.stderr)
    finally:
        if os.path.exists(tmp_docx):
            if pdf_generated_successfully:
                try: os.remove(tmp_docx)
                except OSError as e_del: print(f"WARN (CR): Could not remove temp DOCX {tmp_docx}. {e_del}", file=sys.stderr)
            else:
                print(f"INFO (CR): Temp DOCX {tmp_docx} for {roll_no} KEPT for debugging.", file=sys.stderr)

    return (pdf_name, csv_name) if pdf_generated_successfully else (None, csv_name)


def build_class_pdf(df: pd.DataFrame, image_path: str, exam_details: Dict[str, Any]) -> Tuple[Union[str, None], str]:
    if df.empty:
        print("WARNING (CR): DataFrame for combined class PDF empty. CSV empty, PDF not generated.", file=sys.stderr)
        csv_name = os.path.join(OUTPUT_DIR_COMBINED, "class_marksheet_combined_empty.csv")
        pd.DataFrame().to_csv(csv_name, index=False)
        return None, csv_name

    course_arg = exam_details.get("course_arg", "COURSE") 
    subject_code_arg = exam_details.get("subject_code_arg", "SUBJECT_CODE")
    exam_type_arg = exam_details.get("exam_type_arg", "EXAM_TYPE")
    section_type_arg = exam_details.get("section_type_arg", "SECTION")
    subject_name = exam_details.get("subject", "N/A_SubjectName") 
    year_arg_str = exam_details.get("year", "N/A_Year")
    semester_arg_str = exam_details.get("semester", "N/A_Semester")

    base_filename = f"CLASS_COMBINED_{course_arg}_{subject_code_arg}_{exam_type_arg}_{section_type_arg}_marksheet"
    csv_name = os.path.join(OUTPUT_DIR_COMBINED, f"{base_filename}.csv")
    tmp_docx = os.path.join(OUTPUT_DIR_COMBINED, f"tmp_{base_filename}.docx")
    pdf_name = os.path.join(OUTPUT_DIR_COMBINED, f"{base_filename}.pdf")

    df['max_marks'] = pd.to_numeric(df['max_marks'], errors='coerce').fillna(0).astype(int)
    df['score'] = pd.to_numeric(df['score'], errors='coerce').fillna(0).astype(int)
    
    max_marks_series = df[df['question_id'] != 'N/A'].groupby('question_id')['max_marks'].first().fillna(0).astype(int)
    score_pivot = df.pivot_table(index="roll_no", columns="question_id", values="score", aggfunc="first").fillna(0).astype(int)

    if 'N/A' in score_pivot.columns:
        score_pivot = score_pivot.drop(columns=['N/A'], errors='ignore')

    total_max_class = max_marks_series.sum() if not max_marks_series.empty else 0
    
    qids_for_total_sum = [col for col in score_pivot.columns if col in max_marks_series.index]
    score_pivot["Total"] = score_pivot[qids_for_total_sum].sum(axis=1)
    score_pivot["Percentage"] = ((score_pivot["Total"] / total_max_class) * 100).round(2) if total_max_class > 0 else 0.0

    renamed_cols = {qid: f"{qid}({int(marks)})" for qid, marks in max_marks_series.items() if qid in score_pivot.columns}
    score_pivot_renamed = score_pivot.rename(columns=renamed_cols)

    df_for_csv = score_pivot_renamed.copy()
    if df_for_csv.index.name == "roll_no":
         df_for_csv = df_for_csv.reset_index()
    df_for_csv.rename(columns={'roll_no': 'roll no'}, inplace=True)
    cols_order_csv = [col for col in df_for_csv.columns if col not in ['Total', 'Percentage']]
    if 'Total' in df_for_csv.columns: cols_order_csv.append('Total')
    if 'Percentage' in df_for_csv.columns: cols_order_csv.append('Percentage')
    df_for_csv = df_for_csv[cols_order_csv]
    df_for_csv.to_csv(csv_name, index=False)

    doc = Document()
    sections_doc = doc.sections
    for section in sections_doc: 
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # --- New Header Start (Class PDF) ---
    if os.path.exists(image_path):
        try:
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(image_path, width=Inches(4.0)) # Adjusted "good size"
        except Exception as e_img:
            print(f"WARN (CR Class PDF): Logo add error. {e_img}", file=sys.stderr)
    else:
        print(f"WARN (CR Class PDF): Logo missing at '{image_path}'.", file=sys.stderr)
        
    inst_name_p1 = doc.add_paragraph()
    inst_name_p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_inst_name1 = inst_name_p1.add_run("NATIONAL INSTITUTE OF TECHNOLOGY")
    run_inst_name1.font.name = 'Times New Roman'
    run_inst_name1.font.size = Pt(16)
    run_inst_name1.font.bold = True

    inst_name_p2 = doc.add_paragraph()
    inst_name_p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_inst_name2 = inst_name_p2.add_run("TIRUCHIRAPPALLI")
    run_inst_name2.font.name = 'Times New Roman'
    run_inst_name2.font.size = Pt(14)
    run_inst_name2.font.bold = True
    
    doc.add_paragraph() # Spacer paragraph after header
    # --- New Header End (Class PDF) ---
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tr_run = p_title.add_run("COMBINED CLASS MARKSHEET")
    tr_run.font.name = 'Times New Roman'; tr_run.font.size = Pt(16); tr_run.font.bold = True
    
    details_block_title_paragraph = doc.add_paragraph()
    details_block_title_run = details_block_title_paragraph.add_run("Examination Details")
    details_block_title_run.font.name = 'Times New Roman'; details_block_title_run.font.size = Pt(12); details_block_title_run.font.bold = True
    details_block_title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    details_data = [
        ("Program:", course_arg),
        ("Course Name:", subject_name),
        ("Course Code:", subject_code_arg),
        ("Exam Type:", exam_type_arg),
        ("Section:", section_type_arg),
        ("Year:", year_arg_str),
        ("Semester:", semester_arg_str)
    ]
    details_table = doc.add_table(rows=len(details_data), cols=2)
    details_table.style = 'Table Grid'
    for i, (label, value) in enumerate(details_data):
        label_cell_p = details_table.cell(i, 0).paragraphs[0]; label_cell_p.text = ''
        run_label = label_cell_p.add_run(label)
        run_label.font.name = 'Times New Roman'; run_label.font.size = Pt(10); run_label.font.bold = True
        
        value_cell_p = details_table.cell(i, 1).paragraphs[0]; value_cell_p.text = ''
        run_value = value_cell_p.add_run(str(value)) # Ensure value is string
        run_value.font.name = 'Times New Roman'; run_value.font.size = Pt(10); run_value.font.bold = False
    try:
        details_table.columns[0].width = Inches(1.8) 
        details_table.columns[1].width = Inches(4.2) 
    except Exception as e_width_details:
        print(f"WARN (CR Class): Error setting details_table column widths: {e_width_details}", file=sys.stderr)
    doc.add_paragraph() 

    tbl_data_docx = score_pivot_renamed.reset_index()
    tbl_data_docx.rename(columns={'roll_no': 'roll no'}, inplace=True)
    cols_order_docx = [col for col in tbl_data_docx.columns if col not in ['Total', 'Percentage']]
    if 'Total' in tbl_data_docx.columns: cols_order_docx.append('Total')
    if 'Percentage' in tbl_data_docx.columns: cols_order_docx.append('Percentage')
    tbl_data_docx = tbl_data_docx[cols_order_docx]

    if not tbl_data_docx.empty:
        num_cols = len(tbl_data_docx.columns)
        ct = doc.add_table(rows=1, cols=num_cols)
        ct.style = 'Table Grid'
        ct.autofit = False
        try: 
            if num_cols > 0: ct.columns[0].width = Inches(1.3)  
            num_qid_cols = num_cols - 1 - (1 if 'Total' in cols_order_docx else 0) - (1 if 'Percentage' in cols_order_docx else 0)
            qid_col_width = Inches(0.7) if num_qid_cols > 0 else Inches(0.8)
            for i in range(1, 1 + num_qid_cols):
                if i < num_cols: ct.columns[i].width = qid_col_width
            current_col_idx = 1 + num_qid_cols
            if 'Total' in cols_order_docx and current_col_idx < num_cols: ct.columns[current_col_idx].width = Inches(0.8); current_col_idx +=1
            if 'Percentage' in cols_order_docx and current_col_idx < num_cols: ct.columns[current_col_idx].width = Inches(1.0)
        except Exception as e_width_main: print(f"WARN (CR Class): Error setting main table column widths: {e_width_main}", file=sys.stderr)

        hdr_cells = ct.rows[0].cells
        for i, column_name in enumerate(tbl_data_docx.columns):
            cell_p_hdr = hdr_cells[i].paragraphs[0]; cell_p_hdr.text = ''
            run_hdr = cell_p_hdr.add_run(str(column_name))
            run_hdr.font.name = 'Times New Roman'; run_hdr.font.bold = True; run_hdr.font.size = Pt(9)
            cell_p_hdr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for _, data_row_series in tbl_data_docx.iterrows():
            row_cells = ct.add_row().cells
            for i, cell_value in enumerate(data_row_series):
                cell_p_data = row_cells[i].paragraphs[0]; cell_p_data.text = ''
                run_data = cell_p_data.add_run(str(cell_value))
                run_data.font.name = 'Times New Roman'; run_data.font.size = Pt(9)
                cell_p_data.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if i == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        doc.add_paragraph("No class scores to display.")

    pdf_generated_successfully = False
    try:
        doc.save(tmp_docx)
        pdf_generated_successfully = convert_docx_to_pdf_pandoc(tmp_docx, pdf_name)
    finally:
        if os.path.exists(tmp_docx):
            if pdf_generated_successfully:
                try: os.remove(tmp_docx)
                except OSError as e: print(f"WARN (CR Class): Could not remove temp DOCX {tmp_docx}. {e}", file=sys.stderr)
            else:
                print(f"INFO (CR Class): Temp DOCX {tmp_docx} KEPT for debugging.", file=sys.stderr)
    return (pdf_name, csv_name) if pdf_generated_successfully else (None, csv_name)

# ---------------------------------------------------------------------------
# MAIN PROCESSING FUNCTION
# ---------------------------------------------------------------------------
def process_combined_exam_results(
    course_arg: str,
    subject_code_arg: str,
    exam_type_arg: str,
    year_arg: int,
    semester_arg: int,
    section_type_arg: str,
    logo_img_path_param: str = LOGO_IMAGE_PATH
) -> Dict[str, Any]:
    print(
        f"INFO (CombinedResults): Processing: C:{course_arg}, SubCode:{subject_code_arg}, "
        f"Exam:{exam_type_arg}, Yr:{year_arg}, Sem:{semester_arg}, Sec:{section_type_arg}",
        file=sys.stderr
    )

    prof_criteria_query = {
        "course": course_arg, 
        "subjectCode": subject_code_arg,
        "examType": exam_type_arg,
        "year": year_arg,
        "semester": semester_arg,
        "sectionType": section_type_arg
    }
    professor_doc_main = professoruploads_collection.find_one(prof_criteria_query)

    if not professor_doc_main:
        msg = f"ProfessorUpload document not found for criteria: {prof_criteria_query}"
        print(f"ERROR (CombinedResults): {msg}", file=sys.stderr)
        return {"status": "error_professor_data_missing", "message": msg}

    professor_upload_id = professor_doc_main["_id"]
    professor_subject_name = str(professor_doc_main.get("subject", "N/A_SubjectName"))


    if (professor_doc_main.get("combinedResultGenerationStatus") == "completed_success" and
            professor_doc_main.get("combinedClassResultPdfGridFsId")):
        try:
            if fs_class_aggregate_bucket.exists(ObjectId(professor_doc_main["combinedClassResultPdfGridFsId"])):
                msg = "Combined results previously generated and available."
                print(f"INFO (CombinedResults): {msg} GridFS ID: {professor_doc_main['combinedClassResultPdfGridFsId']}", file=sys.stderr)
                return {
                    "status": "success_already_processed",
                    "message": msg,
                    "combinedClassPdfGridFsId": str(professor_doc_main["combinedClassResultPdfGridFsId"]),
                    "combinedClassCsvGridFsId": str(professor_doc_main.get("combinedClassResultCsvGridFsId")),
                    "data": {"professorUploadId": str(professor_upload_id)}
                }
        except Exception as e_fs_check:
            print(f"WARN (CR): Error checking existing GridFS combined file: {e_fs_check}. Regenerating.", file=sys.stderr)

    professoruploads_collection.update_one(
        {"_id": professor_upload_id},
        {"$set": {
            "combinedResultGenerationStatus": "processing_started",
            "combinedResultInitiatedAt": datetime.now(timezone.utc),
            "combinedResultErrorMessage": None
        }}
    )

    professor_questions_list = professor_doc_main.get("processedJSON", [])
    if not isinstance(professor_questions_list, list) or not professor_questions_list:
        msg = f"ProfessorUpload (ID: {professor_upload_id}) 'processedJSON' is invalid or empty."
        print(f"ERROR (CombinedResults): {msg}", file=sys.stderr)
        professoruploads_collection.update_one(
            {"_id": professor_upload_id},
            {"$set": {"combinedResultGenerationStatus": "error_prof_data_incomplete", "combinedResultErrorMessage": msg}}
        )
        return {"status": "error_prof_data_incomplete", "message": msg}

    reference_data_parsed = parse_reference_answers_from_processed_json(professor_questions_list)
    reference_vectors = build_reference_vectors(reference_data_parsed)
    max_marks_map = {q["question_id"]: q["max_marks"] for q in reference_data_parsed.get("questions", [])}
    total_max_marks_from_prof_for_individual = sum(
        int(q.get("marks", 0)) for q in professor_questions_list if isinstance(q, dict) and q.get("marks") is not None
    )

    exam_details_for_individual_pdfs = {
        "examType": exam_type_arg,
        "subjectCode": subject_code_arg,
        "subject": professor_subject_name, 
        "course": course_arg,             
        "sectionType": section_type_arg,
        "year": str(year_arg),
        "semester": str(semester_arg),
        "total_max_marks_from_prof": total_max_marks_from_prof_for_individual,
        "examDetails": professor_doc_main.get("examDetails", {}) 
    }

    all_individual_scored_dfs_for_class_pdf = []
    processed_students_count = 0
    failed_students_processing_count = 0

    student_array_from_prof_doc = professor_doc_main.get("students", [])
    if not isinstance(student_array_from_prof_doc, list) or not student_array_from_prof_doc:
        msg = f"No 'students' array or empty in ProfessorUpload ID: {professor_upload_id}."
        print(f"WARNING (CombinedResults): {msg}", file=sys.stderr)
        professoruploads_collection.update_one(
            {"_id": professor_upload_id},
            {"$set": {"combinedResultGenerationStatus": "error_no_students_in_prof_doc", "combinedResultErrorMessage": msg}}
        )
        return {"status": "error_no_students_in_prof_doc", "message": msg}


    for student_data in student_array_from_prof_doc:
        roll_no = student_data.get("roll_no")
        if not roll_no:
            print(f"WARN (CR): Student entry in {professor_upload_id} missing 'roll_no'. Skipping.", file=sys.stderr)
            failed_students_processing_count += 1
            continue

        print(f"INFO (CombinedResults): Processing student: {roll_no}", file=sys.stderr)
        scored_df_single_student = calculate_similarity_for_student(student_data, reference_vectors, max_marks_map)
        individual_pdf_gridfs_id = None
        individual_csv_abs_path = None
        notes_for_result_doc = "Processed successfully."
        total_obtained = 0
        scores_for_db = []

        if scored_df_single_student.empty:
            print(f"INFO (CR): No scorable answers for {roll_no}. Minimal record being created. Adding placeholders for class PDF.", file=sys.stderr)
            placeholder_rows = []
            qids_for_exam = list(max_marks_map.keys()) if max_marks_map else ["N/A"]
            if not qids_for_exam or (len(qids_for_exam)==1 and qids_for_exam[0]=='N/A' and not max_marks_map) : 
                 print(f"WARN (CR): No questions defined in max_marks_map for student {roll_no} placeholder generation. Using single N/A.", file=sys.stderr)
                 qids_for_exam = ["N/A"]


            for qid in qids_for_exam:
                placeholder_rows.append({
                    "roll_no": roll_no, "question_id": qid,
                    "max_marks": max_marks_map.get(qid, 0),
                    "similarity": 0.0, "score": 0, 
                    "student_answer_summary": "No answer/Not scorable"
                })
            
            df_for_this_student_class_pdf = pd.DataFrame(placeholder_rows)
            all_individual_scored_dfs_for_class_pdf.append(df_for_this_student_class_pdf)
            notes_for_result_doc = "No scorable answers found. Placeholder added to class sheet."
        else:
            individual_pdf_path, individual_csv_path_temp = build_student_pdf(
                scored_df_single_student, roll_no, logo_img_path_param, exam_details_for_individual_pdfs
            )
            if individual_csv_path_temp and os.path.exists(individual_csv_path_temp):
                individual_csv_abs_path = os.path.abspath(individual_csv_path_temp)

            if individual_pdf_path and os.path.exists(individual_pdf_path):
                pdf_fn_indiv = os.path.basename(individual_pdf_path)
                meta_indiv = {
                    "rollNo": roll_no, "courseName": course_arg, "subjectCode": subject_code_arg,
                    "examType": exam_type_arg, "year": year_arg, "semester": semester_arg,
                    "sectionType": section_type_arg, "type": "student_marksheet_individual_from_combined",
                    "professorUploadIdContext": str(professor_upload_id), "generatedAt": datetime.now(timezone.utc)
                }
                del_q_indiv = {"filename": pdf_fn_indiv, "metadata.rollNo": roll_no, "metadata.subjectCode": subject_code_arg, "metadata.examType": exam_type_arg, "metadata.type": meta_indiv["type"]}
                for old_file in fs_individual_results_bucket.find(del_q_indiv): fs_individual_results_bucket.delete(old_file._id)
                with open(individual_pdf_path, "rb") as f_in:
                    individual_pdf_gridfs_id = str(fs_individual_results_bucket.put(f_in, filename=pdf_fn_indiv, contentType='application/pdf', metadata=meta_indiv))
                try:
                    os.remove(individual_pdf_path)
                    if individual_csv_path_temp and os.path.exists(individual_csv_path_temp): os.remove(individual_csv_path_temp) # remove csv only if pdf was made and removed
                except OSError as e: print(f"WARN (CR): Could not remove temp indiv file(s) for {roll_no}. {e}", file=sys.stderr)
            else:
                print(f"WARN (CR): Individual PDF not generated for {roll_no}.", file=sys.stderr)
                notes_for_result_doc = "Individual PDF generation failed."

            total_obtained = int(scored_df_single_student["score"].sum())
            scores_for_db = scored_df_single_student.to_dict(orient="records")
            all_individual_scored_dfs_for_class_pdf.append(scored_df_single_student)

        percentage = round((total_obtained / total_max_marks_from_prof_for_individual) * 100, 2) if total_max_marks_from_prof_for_individual > 0 else 0.0
        student_mongo_id_val = student_data.get("studentMongoId") or student_data.get("_id")
        if student_mongo_id_val and not isinstance(student_mongo_id_val, ObjectId):
            try: student_mongo_id_val = ObjectId(str(student_mongo_id_val))
            except: student_mongo_id_val = None

        criteria_for_results_doc = {
            "courseName": course_arg, "subjectCode": subject_code_arg, "examType": exam_type_arg,
            "year": year_arg, "semester": semester_arg, "sectionType": section_type_arg,
            "examTitleFromProf": f"{exam_details_for_individual_pdfs.get('subject', 'N/A')} - {exam_details_for_individual_pdfs.get('examType', 'N/A')} - Sec {exam_details_for_individual_pdfs.get('sectionType', 'N/A')}"
        }
        individual_result_payload = {
            "studentMongoId": student_mongo_id_val, "professorMongoId": professor_upload_id, "rollNo": roll_no,
            "criteria": criteria_for_results_doc, "scoresPerQuestion": scores_for_db,
            "totalObtainedMarks": total_obtained, "totalMaximumMarks": total_max_marks_from_prof_for_individual,
            "overallPercentage": percentage, "generatedAt": datetime.now(timezone.utc),
            "gridFsPdfId": individual_pdf_gridfs_id, "localCsvPath": individual_csv_abs_path
        }
        if notes_for_result_doc != "Processed successfully.": individual_result_payload["notes"] = notes_for_result_doc
        
        query_criteria_for_results = {"rollNo": roll_no, "professorMongoId": professor_upload_id}
        for k, v in criteria_for_results_doc.items():
            if k != "examTitleFromProf": query_criteria_for_results[f"criteria.{k}"] = v
        results_collection.update_one(query_criteria_for_results, {"$set": individual_result_payload}, upsert=True)
        processed_students_count += 1

    if not all_individual_scored_dfs_for_class_pdf: 
        msg = "No student data (neither scored nor placeholder) available to generate combined class PDF."
        print(f"ERROR (CombinedResults): {msg}", file=sys.stderr)
        professoruploads_collection.update_one(
            {"_id": professor_upload_id},
            {"$set": {"combinedResultGenerationStatus": "error_no_data_for_class_pdf", "combinedResultErrorMessage": msg}}
        )
        return {"status": "error_no_data_for_class_pdf", "message": msg}

    class_df_final = pd.concat(all_individual_scored_dfs_for_class_pdf, ignore_index=True)
    
    exam_details_for_class_pdf = {
        "course_arg": course_arg,           
        "subject_code_arg": subject_code_arg, 
        "exam_type_arg": exam_type_arg,     
        "section_type_arg": section_type_arg, 
        "subject": professor_subject_name,  
        "year": str(year_arg),              
        "semester": str(semester_arg),      
    }

    combined_pdf_path, combined_csv_path = build_class_pdf(class_df_final, logo_img_path_param, exam_details_for_class_pdf)
    combined_pdf_gridfs_id = None
    combined_csv_gridfs_id = None

    if combined_pdf_path and os.path.exists(combined_pdf_path):
        pdf_fn_class = os.path.basename(combined_pdf_path)
        meta_class_pdf = {"professorUploadId": str(professor_upload_id), **prof_criteria_query, "type": "class_marksheet_combined_aggregate_pdf", "generatedByScript": "Combined_Results.py", "generatedAt": datetime.now(timezone.utc)}
        del_q_class_pdf = {"filename": pdf_fn_class, "metadata.professorUploadId": str(professor_upload_id), "metadata.type": meta_class_pdf["type"]}
        for old_file in fs_class_aggregate_bucket.find(del_q_class_pdf): fs_class_aggregate_bucket.delete(old_file._id)
        with open(combined_pdf_path, "rb") as f_class_pdf:
            combined_pdf_gridfs_id = str(fs_class_aggregate_bucket.put(f_class_pdf, filename=pdf_fn_class, contentType='application/pdf', metadata=meta_class_pdf))
        try: os.remove(combined_pdf_path)
        except OSError as e: print(f"WARN (CR): Could not remove temp combined PDF {combined_pdf_path}. {e}", file=sys.stderr)
    else:
        msg = f"Combined Class PDF path not generated or file does not exist: {combined_pdf_path}"
        print(f"ERROR (CombinedResults): {msg}", file=sys.stderr)
        professoruploads_collection.update_one({"_id": professor_upload_id}, {"$set": {"combinedResultGenerationStatus": "error_combined_pdf_generation", "combinedResultErrorMessage": msg}})
        return {"status": "error_combined_pdf_generation", "message": msg}

    if combined_csv_path and os.path.exists(combined_csv_path) and combined_pdf_gridfs_id:
        csv_fn_class = os.path.basename(combined_csv_path)
        meta_class_csv = {"professorUploadId": str(professor_upload_id), **prof_criteria_query, "type": "class_marksheet_combined_aggregate_csv", "generatedByScript": "Combined_Results.py", "generatedAt": datetime.now(timezone.utc), "pdfPairId": combined_pdf_gridfs_id}
        del_q_class_csv = {"filename": csv_fn_class, "metadata.professorUploadId": str(professor_upload_id), "metadata.type": meta_class_csv["type"]}
        for old_file in fs_class_aggregate_bucket.find(del_q_class_csv): fs_class_aggregate_bucket.delete(old_file._id)
        with open(combined_csv_path, "rb") as f_class_csv:
            combined_csv_gridfs_id = str(fs_class_aggregate_bucket.put(f_class_csv, filename=csv_fn_class, contentType='text/csv', metadata=meta_class_csv))
        try: os.remove(combined_csv_path)
        except OSError as e: print(f"WARN (CR): Could not remove temp combined CSV {combined_csv_path}. {e}", file=sys.stderr)
    elif combined_pdf_gridfs_id:
        print(f"WARN (CR): Combined Class CSV not found or not uploaded ({combined_csv_path}), but PDF was successful.", file=sys.stderr)

    final_update_payload = {
        "combinedResultGenerationStatus": "completed_success",
        "combinedClassResultPdfGridFsId": combined_pdf_gridfs_id,
        "combinedClassResultCsvGridFsId": combined_csv_gridfs_id,
        "combinedResultProcessedAt": datetime.now(timezone.utc),
        "combinedResultStudentProcessedCount": processed_students_count,
        "combinedResultStudentFailedOrSkippedCount": failed_students_processing_count,
        "combinedResultErrorMessage": None
    }
    professoruploads_collection.update_one({"_id": professor_upload_id}, {"$set": final_update_payload})
    print(f"INFO (CombinedResults): ProfessorUpload {professor_upload_id} updated with combined result.", file=sys.stderr)

    return {
        "status": "success_combined_generated",
        "message": f"Combined class result processing complete. Processed {processed_students_count} students. Combined PDF and CSV generated.",
        "combinedClassPdfGridFsId": combined_pdf_gridfs_id,
        "combinedClassCsvGridFsId": combined_csv_gridfs_id,
        "processedStudentCount": processed_students_count
    }
# ---------------------------------------------------------------------------
# CLI ENTRY POINT
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Combined Class Results and Individual Student Marksheets.")
    parser.add_argument("--course", required=True, help="Course Name/Program (e.g., MCA)")
    parser.add_argument("--subject_code", required=True, help="Subject Code (e.g., CA712)")
    parser.add_argument("--exam_type", required=True, help="Exam Type (e.g., CT1)")
    parser.add_argument("--year", required=True, type=int, help="Year (e.g., 2025)")
    parser.add_argument("--semester", required=True, type=int, help="Semester (e.g., 2)")
    parser.add_argument("--section", required=True, help="Section (e.g., A)")
    parser.add_argument("--logo_path", default=LOGO_IMAGE_PATH, help="Path to the logo image file.")

    args = parser.parse_args()
    cli_output_result = {}
    try:
        print(f"--- CLI (CombinedResults): Processing for Course: {args.course}, Subject: {args.subject_code}, Exam: {args.exam_type}, Section: {args.section} ---", file=sys.stderr)
        cli_output_result = process_combined_exam_results(
            course_arg=args.course, subject_code_arg=args.subject_code, exam_type_arg=args.exam_type,
            year_arg=args.year, semester_arg=args.semester, section_type_arg=args.section,
            logo_img_path_param=args.logo_path
        )
    except Exception as e:
        cli_output_result = {"status": "error_cli_exception", "message": str(e), "traceback": traceback.format_exc()}
        print(f"CLI (CombinedResults) Unhandled Error: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
    finally:
        if 'client' in globals() and client:
            try: client.close(); print("INFO (CombinedResults CLI): MongoDB connection closed.", file=sys.stderr)
            except Exception as e_close: print(f"WARN (CombinedResults CLI): Error closing MongoDB: {e_close}", file=sys.stderr)
        print(json.dumps(cli_output_result, indent=2, default=str))