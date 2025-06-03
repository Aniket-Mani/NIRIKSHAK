


# Set TOKENIZERS_PARALLELISM to false before importing sentence_transformers
import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import json
import re
import sys
import traceback
import time # For adding small delays
from typing import Dict, List, Union, Any, Tuple
import pandas as pd
import numpy as np
from pymongo import MongoClient, errors as PyMongoErrors
from bson.objectid import ObjectId
import gridfs
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document # For creating DOCX
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx2pdf import convert # We are removing this
import subprocess # <--- ADDED for Pandoc
from dotenv import load_dotenv
from datetime import datetime, timezone
import argparse
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ... (All your existing CONFIGURATION, DATABASE, EMBEDDING, DATA PARSING, SIMILARITY logic remains IDENTICAL) ...
# ... (parse_professor_questions, build_reference_vectors, similarity_dataframe remain IDENTICAL)

# ---------------------------------------------------------------------------
# CONFIGURATION (Assuming this section is correct from your file)
# ---------------------------------------------------------------------------
try:
    # Use __file__ if this script is directly run, otherwise, ensure path is correct
    current_script_path = os.path.dirname(os.path.abspath(__file__))
    env_path = os.path.abspath(os.path.join(current_script_path, '../.env'))
    if os.path.exists(env_path):
        print(f"INFO (MarksheetGen): Loading .env from: {env_path}", file=sys.stderr)
        load_dotenv(dotenv_path=env_path, verbose=True, override=True)
    else:
        print(f"INFO (MarksheetGen): .env file not found at {env_path}. Using defaults or expecting env vars.", file=sys.stderr)
except Exception as e:
    print(f"WARNING (MarksheetGen): Could not load .env file. Error: {e}", file=sys.stderr)

MONGO_CONNECTION_STRING = os.getenv("MONGO_CONNECTION_STRING", "mongodb://localhost:27017/")
DATABASE_NAME = os.getenv("MONGO_DB_NAME", "smart")

LOGO_IMAGE_FILENAME = os.getenv("LOGO_IMAGE_FILENAME", "logo_name.png")
# Ensure PROJECT_ROOT_DIR is robustly defined
try:
    PROJECT_ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
except NameError: # If __file__ is not defined (e.g. interactive)
    PROJECT_ROOT_DIR = os.path.abspath(os.path.join(os.getcwd(), '../..'))


LOGO_IMAGE_PATH = os.path.join(PROJECT_ROOT_DIR, "backend", "logo", LOGO_IMAGE_FILENAME)

if not os.path.exists(LOGO_IMAGE_PATH):
    print(f"WARNING (MarksheetGen): Logo image not found at '{LOGO_IMAGE_PATH}'. Marksheets may not include the logo.", file=sys.stderr)

OUTPUT_DIR_MARKETSHEETS = os.path.join(PROJECT_ROOT_DIR, "backend", "uploads", "generated_marksheets")
GRIDFS_RESULTS_BUCKET_NAME = "results_marksheets"
RESULTS_COLLECTION_NAME = "Results"

# ---------------------------------------------------------------------------
# DATABASE CONNECTIONS (Assuming this section is correct)
# ---------------------------------------------------------------------------
try:
    print(f"INFO (MarksheetGen): Connecting to MongoDB at {MONGO_CONNECTION_STRING}...", file=sys.stderr)
    client = MongoClient(MONGO_CONNECTION_STRING, serverSelectionTimeoutMS=5000)
    client.admin.command('ping')
    db = client[DATABASE_NAME]
    professoruploads_collection = db["professoruploads"]
    studentuploads_collection = db["studentuploads"]
    results_collection = db[RESULTS_COLLECTION_NAME]
    fs_results_bucket = gridfs.GridFS(db, collection=GRIDFS_RESULTS_BUCKET_NAME)
    print(f"INFO (MarksheetGen): Connected to MongoDB: db='{DATABASE_NAME}'", file=sys.stderr)
except PyMongoErrors.ConnectionFailure as e:
    print(f"FATAL (MarksheetGen): Could not connect to MongoDB. Error: {e}", file=sys.stderr)
    sys.exit(1)
except Exception as e:
    print(f"FATAL (MarksheetGen): An unexpected error occurred during MongoDB setup. Error: {e}", file=sys.stderr)
    traceback.print_exc(file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# EMBEDDING MODEL & HELPERS (Assuming this section is correct)
# ---------------------------------------------------------------------------
try:
    print("INFO (MarksheetGen): Loading sentence embedding model 'all-MiniLM-L6-v2'...", file=sys.stderr)
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    print("INFO (MarksheetGen): Sentence embedding model loaded.", file=sys.stderr)
except Exception as e:
    print(f"FATAL (MarksheetGen): Could not load SentenceTransformer model. Error: {e}", file=sys.stderr)
    traceback.print_exc(file=sys.stderr)
    sys.exit(1)

def preprocess(text: str) -> str:
    if not text or not isinstance(text, str): return ""
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def embed(text_to_embed: str) -> Union[np.ndarray, None]:
    processed_text = preprocess(text_to_embed)
    return embedding_model.encode(processed_text, convert_to_numpy=True) if processed_text else None

def cos_sim(v1: Union[np.ndarray, None], v2: Union[np.ndarray, None]) -> float:
    if v1 is None or v2 is None or v1.size == 0 or v2.size == 0 : return 0.0
    v1_r = v1.reshape(1, -1) if v1.ndim == 1 else v1
    v2_r = v2.reshape(1, -1) if v2.ndim == 1 else v2
    if v1_r.shape[1] != v2_r.shape[1]:
        print(f"WARNING (MarksheetGen): Cosine similarity dim mismatch. v1: {v1_r.shape}, v2: {v2_r.shape}", file=sys.stderr)
        return 0.0
    return float(cosine_similarity(v1_r, v2_r)[0][0])

# ---------------------------------------------------------------------------
# DATA PARSING (Professor's Reference Answers)
# ---------------------------------------------------------------------------
def parse_professor_questions(professor_processed_json_list: List[Dict[str, Any]]) -> Dict[str, Any]:
    items = []
    for q_idx, q_data in enumerate(professor_processed_json_list):
        question_id_from_prof = q_data.get('questionNo', f"q{q_idx+1}")
        temp_id = str(question_id_from_prof).strip()
        match = re.match(r"^[Qq]([0-9]+[a-zA-Z]?)$", temp_id)
        if match:
            qid_normalized = match.group(1)
        else:
            qid_normalized = temp_id.replace(" ", "").replace(".", "_")
        
        ref_answers_list = q_data.get("Answers", [])
        # Ensure padded_refs always takes string values, defaulting to empty string
        actual_refs = [str(ans) if ans is not None else "" for ans in ref_answers_list]
        padded_refs = (actual_refs + ["", "", ""])[:3]

        items.append({
            "question_id": qid_normalized,
            "max_marks"  : int(q_data.get("marks", 0)),
            "answer1"    : padded_refs[0],
            "answer2"    : padded_refs[1],
            "answer3"    : padded_refs[2],
            "question_text": str(q_data.get("questionText",""))
        })
    return {"questions": items}

def build_reference_vectors(parsed_prof_questions: Dict[str, Any]) -> Dict[str, List[Union[np.ndarray, None]]]:
    vector_cache = {}
    for q in parsed_prof_questions.get("questions", []):
        vector_cache[q["question_id"]] = [embed(q["answer1"]), embed(q["answer2"]), embed(q["answer3"])]
    return vector_cache

# ---------------------------------------------------------------------------
# SIMILARITY DATAFRAME BUILDER
# ---------------------------------------------------------------------------
def similarity_dataframe(
    student_doc: Dict[str, Any],
    ref_vecs: Dict[str, List[Union[np.ndarray, None]]],
    max_marks_map: Dict[str, int]
) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    roll = student_doc.get("username") # Assuming 'username' is the roll number key
    if not roll:
        print(f"WARNING (MarksheetGen): Student document missing 'username' (roll_no): {student_doc.get('_id')}", file=sys.stderr)
        return pd.DataFrame()

    student_answers_list = student_doc.get("extractedAnswer", {}).get("answers", [])
    for ans_data in student_answers_list:
        if not isinstance(ans_data, dict): continue
        qid_from_student = ans_data.get("question_id")
        # Normalize QID consistently
        qid_normalized = str(qid_from_student).strip().replace(" ", "").replace(".", "_") if qid_from_student is not None else None
        stu_answer_text = ans_data.get("answer_text")
        if not qid_normalized or stu_answer_text is None: continue

        stu_vec = embed(str(stu_answer_text)) # Ensure text is string
        similarity = 0.0
        if stu_vec is not None:
            ref_q_vecs = ref_vecs.get(qid_normalized, [])
            if not ref_q_vecs and qid_normalized in max_marks_map: # Check if QID was expected
                print(f"WARNING (MarksheetGen): No reference vectors found for QID '{qid_normalized}' which has max_marks. Similarity will be 0.", file=sys.stderr)
            
            sim_scores = [cos_sim(stu_vec, ref_v) for ref_v in ref_q_vecs if ref_v is not None]
            similarity = max(sim_scores) if sim_scores else 0.0

        question_max_marks = max_marks_map.get(qid_normalized, 0)
        if qid_normalized not in max_marks_map:
             # This might be okay if student answered a question not in prof's list, or vice-versa.
             print(f"INFO (MarksheetGen): Max marks not found for QID '{qid_normalized}' (student answer). Max marks will be 0.", file=sys.stderr)

        rows.append({
            "roll_no"    : roll,
            "question_id": qid_normalized,
            "max_marks"  : question_max_marks,
            "similarity" : round(similarity, 3),
            "student_answer_summary": str(stu_answer_text)[:100] + "..." if stu_answer_text and len(str(stu_answer_text)) > 100 else str(stu_answer_text)
        })
    df = pd.DataFrame(rows)
    if df.empty:
        print(f"INFO (MarksheetGen): Similarity dataframe is empty for roll {roll}.", file=sys.stderr)
        # Ensure it still has the necessary columns for concatenation later, if applicable
        return pd.DataFrame(columns=["roll_no", "question_id", "max_marks", "similarity", "student_answer_summary", "score"])


    def score_rule(sim: float, max_m: int) -> int:
        # Ensure sim and max_m are numeric
        sim = float(sim) if sim is not None else 0.0
        max_m = int(max_m) if max_m is not None else 0
        
        if sim >= 0.90: pct = 1.00
        elif sim >= 0.80: pct = 0.90
        elif sim >= 0.70: pct = 0.75
        elif sim >= 0.60: pct = 0.65
        elif sim >= 0.50: pct = 0.60
        elif sim >= 0.45: pct = 0.50
        else: pct = 0.0
        return int(round(pct * max_m))

    df["score"] = df.apply(lambda r: score_rule(r["similarity"], r["max_marks"]), axis=1)
    return df


# ---------------------------------------------------------------------------
# DOCX / PDF BUILDERS
# ---------------------------------------------------------------------------
def build_student_pdf_with_pandoc(df_student_scores: pd.DataFrame, roll_no: str, logo_image_path_param: str, exam_details_for_pdf_gen: Dict[str, Any]) -> Tuple[Union[str,None], Union[str,None]]:
    os.makedirs(OUTPUT_DIR_MARKETSHEETS, exist_ok=True)
    exam_type_short = str(exam_details_for_pdf_gen.get("examType", "Exam")).replace(" ", "_")[:20]
    subject_code_short = str(exam_details_for_pdf_gen.get("subjectCode", "UnknownSub"))
    course_short = str(exam_details_for_pdf_gen.get("course", "UnknownCourse")).replace(" ","_")[:15]
    section_short = str(exam_details_for_pdf_gen.get("sectionType", "AllSections"))
    base_filename = f"{roll_no}_{course_short}_{subject_code_short}_{exam_type_short}_{section_short}_marksheet"
    
    csv_name = os.path.join(OUTPUT_DIR_MARKETSHEETS, f"{base_filename}.csv")
    pdf_name = os.path.join(OUTPUT_DIR_MARKETSHEETS, f"{base_filename}.pdf")
    tmp_docx = os.path.join(OUTPUT_DIR_MARKETSHEETS, f"tmp_{base_filename}.docx")

    # Ensure numeric types for calculations
    df_student_scores['max_marks'] = pd.to_numeric(df_student_scores['max_marks'], errors='coerce').fillna(0).astype(int)
    df_student_scores['score'] = pd.to_numeric(df_student_scores['score'], errors='coerce').fillna(0).astype(int)
    
    df_student_scores["percentage"] = df_student_scores.apply(
        lambda row: round((row["score"] / row["max_marks"]) * 100, 2) if row["max_marks"] > 0 else 0.0, axis=1
    )
    
    total_score = df_student_scores["score"].sum() if not df_student_scores.empty else 0
    # Get total_max_marks from professor_doc if available, otherwise sum from df
    total_max_from_prof = exam_details_for_pdf_gen.get("total_max_marks_from_prof", 0)
    total_max_from_df = df_student_scores["max_marks"].sum() if not df_student_scores.empty else 0

    if not df_student_scores.empty and total_max_from_prof > 0 and total_max_from_df != total_max_from_prof:
        print(f"WARNING (MarksheetGen): Discrepancy in max marks sum. DF sum: {total_max_from_df}, Prof sum: {total_max_from_prof}. Using professor's total for overall percentage if available.", file=sys.stderr)
        total_max = total_max_from_prof
    elif total_max_from_prof > 0: # df might be empty but prof total is known
        total_max = total_max_from_prof
    else: # Fallback to df sum if prof total not available
        total_max = total_max_from_df

    total_pct = round(total_score / total_max * 100, 2) if total_max > 0 else 0.0

    summary_data = [{"question_id": "Total", "max_marks": total_max, "score": total_score, "percentage": total_pct}]
    summary_df = pd.DataFrame(summary_data)
    
    cols_for_sheet = ["question_id", "max_marks", "score", "percentage"]
    if not df_student_scores.empty:
        # Ensure all data is string for display in table, to prevent type issues with concat
        df_display = df_student_scores[cols_for_sheet].astype(str)
        summary_display_df = summary_df[cols_for_sheet].astype(str)
        sheet = pd.concat([df_display, summary_display_df], ignore_index=True)
    else:
        sheet = summary_df[cols_for_sheet].astype(str)


    try:
        # Save the logical sheet (with numbers) to CSV
        logical_sheet_for_csv = pd.concat([df_student_scores[cols_for_sheet], summary_df[cols_for_sheet]], ignore_index=True) if not df_student_scores.empty else summary_df[cols_for_sheet]
        logical_sheet_for_csv.to_csv(csv_name, index=False)
        print(f"INFO (MarksheetGen): Student scores CSV saved to {csv_name}", file=sys.stderr)
    except Exception as e:
        print(f"ERROR (MarksheetGen): Could not save CSV {csv_name}. Error: {e}", file=sys.stderr)

    # --- DOCX Generation Start (Identical to your original) ---
    doc = Document()
    # Add logo if it exists
  # Add logo if it exists, using a table for better centering control
    if os.path.exists(logo_image_path_param):
        try:
        # Add a centered paragraph at the beginning
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the image directly into the paragraph
            run = paragraph.add_run()
            run.add_picture(logo_image_path_param, width=Inches(3.0))  # Adjust width as needed

        except Exception as e_img:
            print(f"WARNING (MarksheetGen): Could not add logo image '{logo_image_path_param}'. Error: {e_img}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
    else:
        print(f"WARNING (MarksheetGen): Logo image not found at '{logo_image_path_param}' for student {roll_no}.", file=sys.stderr)

    # Header
    head = doc.add_paragraph()
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    exam_title_display = f"{str(exam_details_for_pdf_gen.get('course', 'N/A'))} - {str(exam_details_for_pdf_gen.get('subject', 'N/A'))} ({str(exam_details_for_pdf_gen.get('subjectCode', 'N/A'))}) - {str(exam_details_for_pdf_gen.get('examType', 'Exam'))} - Section {str(exam_details_for_pdf_gen.get('sectionType', 'N/A'))}"
    hd_run = head.add_run(exam_title_display.upper())
    hd_run.font.size = Pt(14)
    hd_run.font.bold = True
    doc.add_paragraph() # Spacer

    # Title "Marksheet"
    ttl = doc.add_paragraph()
    ttl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tr_run = ttl.add_run("Marksheet")
    tr_run.font.size = Pt(16)
    tr_run.font.bold = True

    # Student Details Table
    details_table = doc.add_table(rows=5, cols=2)
    details_table.style = 'Table Grid'
    details_table.rows[0].cells[0].text = "Roll No:"
    details_table.rows[0].cells[1].text = str(roll_no)
    details_table.rows[1].cells[0].text = "Course Code:"
    details_table.rows[1].cells[1].text = str(exam_details_for_pdf_gen.get("subjectCode", ""))
    details_table.rows[2].cells[0].text = "Course Name:"
    details_table.rows[2].cells[1].text = str(exam_details_for_pdf_gen.get("subject", ""))
    details_table.rows[3].cells[0].text = "Exam Type:"
    details_table.rows[3].cells[1].text = str(exam_details_for_pdf_gen.get("examType", ""))
    details_table.rows[4].cells[0].text = "Section:"
    details_table.rows[4].cells[1].text = str(exam_details_for_pdf_gen.get("sectionType", ""))
    doc.add_paragraph() # Spacer

    # Question-wise Marks Table
    qw = doc.add_paragraph()
    qw_run = qw.add_run("Question-wise Marks:")
    qw_run.font.bold = True
    qw_run.font.size = Pt(12)

    if not sheet.empty:
        marks_table = doc.add_table(rows=1, cols=len(sheet.columns))
        marks_table.style = 'Table Grid'
        # Header row
        for i, column_name in enumerate(sheet.columns):
            marks_table.cell(0, i).text = str(column_name)
        # Data rows
        for r_idx, r_data in sheet.iterrows():
            row_cells = marks_table.add_row().cells
            for c_idx, cell_value in enumerate(r_data):
                row_cells[c_idx].text = str(cell_value)
    else:
        doc.add_paragraph("No scores to display.")
    # --- DOCX Generation End ---

    pdf_generated_successfully = False
    try:
        doc.save(tmp_docx)
        print(f"INFO (MarksheetGen): Temporary DOCX saved to {tmp_docx}", file=sys.stderr)
        
        # --- PANDOC CONVERSION ---
        print(f"INFO (MarksheetGen): Attempting to convert DOCX '{tmp_docx}' to PDF '{pdf_name}' using Pandoc...", file=sys.stderr)
        pandoc_command = [
            'pandoc', tmp_docx,
            '-o', pdf_name,
            # '--pdf-engine=xelatex', # Or other engines like weasyprint, if you have specific needs and they are installed
            # You can add more pandoc options here if needed, e.g., for margins, fonts
            # For example, to specify a LaTeX engine (if you have LaTeX installed):
            # '--pdf-engine=pdflatex' # or xelatex, lualatex
        ]
        # To handle potential issues with default fonts or characters, especially if not using LaTeX:
        # pandoc_command.extend(['--variable', 'mainfont="DejaVu Sans"']) # Example: using a common font

        process = subprocess.run(pandoc_command, capture_output=True, text=True, check=False)

        if process.returncode == 0:
            if os.path.exists(pdf_name) and os.path.getsize(pdf_name) > 0:
                pdf_generated_successfully = True
                print(f"INFO (MarksheetGen): PDF for student {roll_no} created successfully using Pandoc at {pdf_name}", file=sys.stderr)
            else:
                print(f"ERROR (MarksheetGen): Pandoc command succeeded (return code 0) but PDF '{pdf_name}' was not created or is empty.", file=sys.stderr)
                print(f"Pandoc stdout:\n{process.stdout}", file=sys.stderr)
                print(f"Pandoc stderr:\n{process.stderr}", file=sys.stderr)
        else:
            print(f"ERROR (MarksheetGen): Pandoc conversion failed for student {roll_no}. Return code: {process.returncode}", file=sys.stderr)
            print(f"Pandoc stdout:\n{process.stdout}", file=sys.stderr)
            print(f"Pandoc stderr:\n{process.stderr}", file=sys.stderr)
            # Consider keeping tmp_docx if pandoc fails, for debugging the docx
            # return tmp_docx, csv_name # Or handle as error

    except FileNotFoundError:
        print(f"ERROR (MarksheetGen): Pandoc command not found. Please ensure Pandoc is installed and in your system's PATH.", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
    except Exception as e_conv:
        print(f"ERROR (MarksheetGen): DOCX to PDF conversion using Pandoc failed for student {roll_no}. Error: {e_conv}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
    finally:
        # Always try to remove the temp DOCX if it exists, unless Pandoc failed and you want to inspect it
        if os.path.exists(tmp_docx):
            if pdf_generated_successfully : # Only remove if PDF was made
                try:
                    os.remove(tmp_docx)
                    print(f"INFO (MarksheetGen): Temp DOCX file {tmp_docx} removed.", file=sys.stderr)
                except OSError as e_del:
                    print(f"WARNING (MarksheetGen): Could not remove temp file {tmp_docx}. Error: {e_del}", file=sys.stderr)
            else:
                print(f"INFO (MarksheetGen): Temp DOCX file {tmp_docx} was KEPT for debugging due to Pandoc failure.", file=sys.stderr)


    if pdf_generated_successfully:
        return pdf_name, csv_name
    else:
        print(f"WARNING (MarksheetGen): PDF generation ultimately failed or file not confirmed for {roll_no}. Returning None for PDF path.", file=sys.stderr)
        # If Pandoc fails, you might want to return the path to the DOCX for manual inspection/conversion
        # or just None for PDF path to indicate failure.
        return None, csv_name # Returning None for PDF if Pandoc failed

# Alias the new function to the old name if the rest of the script uses `build_student_pdf`
build_student_pdf = build_student_pdf_with_pandoc


# ... (Rest of your script: generate_student_result_service and __main__ block remain IDENTICAL) ...
# ... They will now use the aliased `build_student_pdf` which internally calls `build_student_pdf_with_pandoc`

# ---------------------------------------------------------------------------
# MAIN SERVICE FUNCTIONS
# ---------------------------------------------------------------------------
def generate_student_result_service(
    student_roll_no_arg: str,
    course_arg: str,
    subject_code_arg: str,
    exam_type_arg: str,
    year_arg: int,
    semester_arg: int,
    section_type_arg: str,
    logo_img_path_param: str = LOGO_IMAGE_PATH # Default to global LOGO_IMAGE_PATH
) -> Dict[str, Any]:
    print(f"INFO (MarksheetGen): Service call - RollNo: {student_roll_no_arg}, Course: {course_arg}, SubjectCode: {subject_code_arg}, ExamType: {exam_type_arg}, Year: {year_arg}, Sem: {semester_arg}, Section: {section_type_arg}", file=sys.stderr)

    # Define criteria for querying professor and student data
    common_criteria_for_prof = {
        "course": course_arg, 
        "subjectCode": subject_code_arg,
        "examType": exam_type_arg, 
        "year": year_arg, # Assuming year is stored as number in prof DB
        "semester": semester_arg, # Assuming sem is stored as number
        "sectionType": section_type_arg
    }
    # Student uploads usually have 'username' for roll number
    student_query_criteria = {
        "username": student_roll_no_arg,
        "course": course_arg, 
        "subjectCode": subject_code_arg,
        "examType": exam_type_arg,
        # Assuming year and semester are stored as numbers in studentuploads too, adjust if strings
        "year": year_arg, 
        "semester": semester_arg,
        "sectionType": section_type_arg 
    }
    
    # Criteria for checking existing result in 'Results' collection
    result_query_criteria = {
        "rollNo": student_roll_no_arg,
        "criteria.courseName": course_arg,
        "criteria.subjectCode": subject_code_arg,
        "criteria.examType": exam_type_arg,
        "criteria.year": year_arg,
        "criteria.semester": semester_arg,
        "criteria.sectionType": section_type_arg
    }
    result_query_criteria_fallback = { # If sectionType is missing in an old record
        "rollNo": student_roll_no_arg, "criteria.courseName": course_arg,
        "criteria.subjectCode": subject_code_arg, "criteria.examType": exam_type_arg,
        "criteria.year": year_arg, "criteria.semester": semester_arg,
        # Not including "criteria.sectionType" here
    }

    # 1. Check for cached result
    existing_result = results_collection.find_one(result_query_criteria)
    if not existing_result : # Try fallback if main query failed
        print(f"INFO (MarksheetGen): No cached result with full criteria, trying fallback for {student_roll_no_arg}", file=sys.stderr)
        existing_result = results_collection.find_one(result_query_criteria_fallback)


    if existing_result and existing_result.get("gridFsPdfId"):
        try:
            if fs_results_bucket.exists(ObjectId(existing_result["gridFsPdfId"])):
                 print(f"INFO (MarksheetGen): Cached result for {student_roll_no_arg} found. GridFS ID: {existing_result['gridFsPdfId']}", file=sys.stderr)
                 return {"status": "success_cached", "message": "Result previously generated and found.", "data": existing_result}
            else: # Record exists but GridFS file is missing
                print(f"WARNING (MarksheetGen): Cached result for {student_roll_no_arg} (ID: {existing_result['_id']}) points to missing GridFS file ID {existing_result['gridFsPdfId']}. Regenerating.", file=sys.stderr)
        except Exception as e_gridfs_check: # Invalid ObjectId or other GridFS error
             print(f"WARNING (MarksheetGen): Error checking GridFS for {existing_result.get('gridFsPdfId', 'N/A')}: {e_gridfs_check}. Regenerating result for {student_roll_no_arg}.", file=sys.stderr)

    # 2. Fetch Student Data
    print(f"DEBUG (MarksheetGen): Querying 'studentuploads' with: {student_query_criteria}", file=sys.stderr)
    student_doc = studentuploads_collection.find_one(student_query_criteria)
    if not student_doc:
        # Try a slightly more relaxed query for student doc if sectionType might be missing or different
        relaxed_student_query = student_query_criteria.copy()
        del relaxed_student_query["sectionType"] # Example: try without section
        print(f"DEBUG (MarksheetGen): Retrying 'studentuploads' with relaxed query: {relaxed_student_query}", file=sys.stderr)
        student_doc = studentuploads_collection.find_one(relaxed_student_query)
        if not student_doc:
            msg = f"Student data not found for RollNo: {student_roll_no_arg} with criteria: {student_query_criteria} (and relaxed)."
            print(f"ERROR (MarksheetGen): {msg}", file=sys.stderr)
            return {"status": "error_student_data_missing", "message": msg}

    student_answers_list = student_doc.get("extractedAnswer", {}).get("answers", [])
    if not student_answers_list: # Check if list exists and is not empty
        msg = f"Student {student_roll_no_arg} (Doc ID: {student_doc.get('_id')}) has no extracted answers."
        print(f"ERROR (MarksheetGen): {msg}", file=sys.stderr)
        return {"status": "error_student_extraction_incomplete", "message": msg}

    # 3. Fetch Professor Data
    print(f"DEBUG (MarksheetGen): Querying 'professoruploads' with: {common_criteria_for_prof}", file=sys.stderr)
    professor_doc = professoruploads_collection.find_one(common_criteria_for_prof)
    if not professor_doc:
        # Try a slightly more relaxed query for professor doc if sectionType might be missing or different
        relaxed_prof_query = common_criteria_for_prof.copy()
        del relaxed_prof_query["sectionType"] # Example: try without section
        print(f"DEBUG (MarksheetGen): Retrying 'professoruploads' with relaxed query: {relaxed_prof_query}", file=sys.stderr)
        professor_doc = professoruploads_collection.find_one(relaxed_prof_query)
        if not professor_doc:
            msg = f"Professor's reference data not found for criteria: {common_criteria_for_prof} (and relaxed)."
            print(f"ERROR (MarksheetGen): {msg}", file=sys.stderr)
            return {"status": "error_professor_data_missing", "message": msg}

    professor_questions_list = professor_doc.get("processedJSON", []) # Assuming this is the key
    if not isinstance(professor_questions_list, list) or not professor_questions_list:
        msg = f"Professor (Doc ID: {professor_doc.get('_id')}) 'processedJSON' (questions list) is invalid or empty."
        print(f"ERROR (MarksheetGen): {msg}", file=sys.stderr)
        return {"status": "error_professor_data_incomplete", "message": msg}

    # Calculate total_max_marks from professor's question list
    total_max_marks_from_prof = sum(int(q_data.get("marks", 0)) for q_data in professor_questions_list if isinstance(q_data, dict))
    
    # Prepare exam details for PDF generation, ensuring all values are strings for safety
    exam_details_for_pdf_generation = {
        "examType": str(professor_doc.get("examType", exam_type_arg)),
        "subjectCode": str(professor_doc.get("subjectCode", subject_code_arg)),
        "subject": str(professor_doc.get("subject", "N/A")),
        "course": str(professor_doc.get("course", course_arg)),
        "sectionType": str(professor_doc.get("sectionType", section_type_arg)), # Use section from prof_doc if available
        "year": str(professor_doc.get("year", year_arg)),
        "semester": str(professor_doc.get("semester", semester_arg)),
        "total_max_marks_from_prof": total_max_marks_from_prof
    }

    # 4. Parse, Build Vectors, Calculate Scores
    print(f"INFO (MarksheetGen): Parsing professor's questions for {subject_code_arg}...", file=sys.stderr)
    reference_data_parsed = parse_professor_questions(professor_questions_list)
    print(f"INFO (MarksheetGen): Building reference vectors for {len(reference_data_parsed.get('questions',[]))} questions...", file=sys.stderr)
    reference_vectors = build_reference_vectors(reference_data_parsed)
    max_marks_map = {q["question_id"]: q["max_marks"] for q in reference_data_parsed.get("questions", [])}

    print(f"INFO (MarksheetGen): Calculating scores for student {student_roll_no_arg}...", file=sys.stderr)
    scored_df = similarity_dataframe(student_doc, reference_vectors, max_marks_map) # student_doc contains 'username' and 'extractedAnswer'

    # 5. Generate PDF using the new Pandoc function (aliased to build_student_pdf)
    print(f"INFO (MarksheetGen): Building PDF for student {student_roll_no_arg} using Pandoc...", file=sys.stderr)
    pdf_file_path, csv_file_path = build_student_pdf(scored_df, student_roll_no_arg, logo_img_path_param, exam_details_for_pdf_generation)

    # 6. Upload to GridFS
    gridfs_file_id_str = None
    # Check if pdf_file_path is not None and points to an actual PDF file
    if pdf_file_path and pdf_file_path.lower().endswith(".pdf") and os.path.exists(pdf_file_path) and os.path.getsize(pdf_file_path) > 0:
        pdf_filename_on_disk = os.path.basename(pdf_file_path)
        # Define metadata for GridFS to make search/delete more specific
        gridfs_metadata = {
            "rollNo": student_roll_no_arg, "courseName": course_arg,
            "subjectCode": subject_code_arg, "examType": exam_type_arg,
            "year": year_arg, "semester": semester_arg,
            "sectionType": exam_details_for_pdf_generation["sectionType"], # Use section from prof doc for consistency
            "type": "student_marksheet",
            "generatedAt": datetime.now(timezone.utc) # Use timezone-aware datetime
        }
        # Delete existing GridFS files matching this specific metadata to prevent duplicates from reruns
        delete_query_fs = {"filename": pdf_filename_on_disk, "metadata.rollNo": student_roll_no_arg, "metadata.subjectCode": subject_code_arg, "metadata.examType": exam_type_arg} # Simplified delete query
        for old_file in fs_results_bucket.find(delete_query_fs):
            try:
                fs_results_bucket.delete(old_file._id)
                print(f"INFO (MarksheetGen): Deleted old GridFS file: {old_file.filename} ID: {old_file._id}", file=sys.stderr)
            except Exception as e_del_fs:
                print(f"WARNING (MarksheetGen): Could not delete old GridFS file {old_file.filename}. Error: {e_del_fs}", file=sys.stderr)

        with open(pdf_file_path, "rb") as f_pdf:
            gridfs_id = fs_results_bucket.put(
                f_pdf, filename=pdf_filename_on_disk, contentType='application/pdf',
                metadata=gridfs_metadata
            )
            gridfs_file_id_str = str(gridfs_id)
        print(f"INFO (MarksheetGen): Marksheet '{pdf_filename_on_disk}' uploaded to GridFS. ID: {gridfs_file_id_str}", file=sys.stderr)
    elif pdf_file_path and not (pdf_file_path.lower().endswith(".pdf") and os.path.exists(pdf_file_path) and os.path.getsize(pdf_file_path) > 0) :
        print(f"ERROR (MarksheetGen): PDF path '{pdf_file_path}' was returned, but it's not a valid PDF file or is empty. Not uploaded.", file=sys.stderr)
    elif not pdf_file_path: # pdf_file_path is None
         print(f"INFO (MarksheetGen): No PDF generated (path is None) for {student_roll_no_arg}. Nothing to upload to GridFS.", file=sys.stderr)


    # 7. Save result summary to 'results' collection
    # Recalculate total_obtained and total_max based on scored_df to be sure
    if not scored_df.empty:
        total_obtained = int(scored_df["score"].sum())
        # For total_max, prioritize the sum from professor's data if available and consistent
        # This was already handled by total_max in build_student_pdf, using total_max_marks_from_prof there
        # For consistency, we can re-fetch it or use the value from exam_details_for_pdf_generation
        total_max_for_result = exam_details_for_pdf_generation["total_max_marks_from_prof"]
        if total_max_for_result == 0 and not scored_df.empty: # Fallback if prof marks somehow 0
            total_max_for_result = int(scored_df["max_marks"].sum())

    else: # scored_df is empty
        total_obtained = 0
        total_max_for_result = exam_details_for_pdf_generation["total_max_marks_from_prof"]


    percentage = round((total_obtained / total_max_for_result) * 100, 2) if total_max_for_result > 0 else 0.0
    scores_to_save = scored_df.to_dict(orient="records") if not scored_df.empty else []

    result_document_payload = {
        "studentMongoId": student_doc["_id"], 
        "professorMongoId": professor_doc["_id"],
        "rollNo": student_roll_no_arg,
        "criteria": { # Store the criteria used for this result generation
            "courseName": course_arg, 
            "subjectCode": subject_code_arg,
            "examType": exam_type_arg, 
            "year": year_arg, 
            "semester": semester_arg,
            "sectionType": exam_details_for_pdf_generation["sectionType"], # Use section from prof_doc
            "examTitleFromProf": f"{exam_details_for_pdf_generation.get('subject', '')} - {exam_details_for_pdf_generation.get('examType', '')} - Sec {exam_details_for_pdf_generation.get('sectionType','')}"
        },
        "scoresPerQuestion": scores_to_save,
        "totalObtainedMarks": total_obtained, 
        "totalMaximumMarks": total_max_for_result,
        "overallPercentage": percentage, 
        "generatedAt": datetime.now(timezone.utc), # Use timezone-aware UTC datetime
        "gridFsPdfId": gridfs_file_id_str, # This will be None if PDF upload failed
        "localCsvPath": os.path.abspath(csv_file_path) if csv_file_path and os.path.exists(csv_file_path) else None
    }

    try:
        # Use the same result_query_criteria for update_one to ensure consistency
        results_collection.update_one(result_query_criteria, {"$set": result_document_payload}, upsert=True)
        print(f"INFO (MarksheetGen): Result for {student_roll_no_arg} saved/updated in '{RESULTS_COLLECTION_NAME}'.", file=sys.stderr)
    except Exception as e_db_res:
        print(f"ERROR (MarksheetGen): Failed to save result to DB for {student_roll_no_arg}. Error: {e_db_res}", file=sys.stderr)
        # Even if DB save fails, if PDF was made and uploaded, that's partial success
        return {
            "status": "success_generated_db_error",
            "message": f"Marksheet PDF may have been generated for {student_roll_no_arg} (GridFS ID: {gridfs_file_id_str}), but saving result summary to DB failed.",
            "rollNo": student_roll_no_arg, "gridFsPdfId": gridfs_file_id_str, "data": result_document_payload
        }

    # Determine final status based on PDF generation success
    final_status = "success_generated_new"
    final_message = f"Marksheet processing complete for student {student_roll_no_arg}."
    if not gridfs_file_id_str : # This means PDF was not generated or not uploaded
        final_status = "success_scores_calculated_pdf_error" if not scored_df.empty else "success_no_scores_to_process_pdf_error"
        final_message += " However, PDF could not be generated or stored in GridFS."
    elif not scores_to_save: # PDF made, but no scores (e.g., empty student answers)
         final_status = "success_pdf_generated_no_scores"
         final_message += " PDF generated, but no specific question scores were calculated (e.g., student answers were empty)."

    return {
        "status": final_status, "message": final_message, "rollNo": student_roll_no_arg,
        "gridFsPdfId": gridfs_file_id_str, "data": result_document_payload
    }

# ---------------------------------------------------------------------------
# CLI ENTRY POINT (Assuming this section is correct from your file)
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate student marksheet or class marksheet.")
    parser.add_argument("--roll_no", help="Student's Roll Number (for individual marksheet)")
    parser.add_argument("--course", required=True, help="Course Name (e.g., MCA)")
    parser.add_argument("--subject_code", required=True, help="Subject Code (e.g., CA721)")
    parser.add_argument("--exam_type", required=True, help="Exam Type (e.g., CT1)")
    parser.add_argument("--year", required=True, type=int, help="Year (e.g., 2025)")
    parser.add_argument("--semester", required=True, type=int, help="Semester (e.g., 3)")
    parser.add_argument("--section", required=True, help="Section (e.g., A or AllSections)") # Ensure this matches DB
    parser.add_argument("--mode", choices=['student', 'class'], default='student', help="Generate for 'student' or 'class'")

    args = parser.parse_args()
    output_result = {}
    try:
        if args.mode == 'student':
            if not args.roll_no:
                raise ValueError("--roll_no is required for student mode.")
            print(f"--- CLI: Generating Student Marksheet for RollNo: {args.roll_no} ---", file=sys.stderr)
            output_result = generate_student_result_service(
                student_roll_no_arg=args.roll_no,
                course_arg=args.course,
                subject_code_arg=args.subject_code,
                exam_type_arg=args.exam_type,
                year_arg=args.year,
                semester_arg=args.semester,
                section_type_arg=args.section, # Passed to service
                logo_img_path_param=LOGO_IMAGE_PATH
            )
        elif args.mode == 'class':
            # This part would need significant new logic for class-wide processing
            print(f"--- CLI: Generating Class Marksheet for Course: {args.course}, Subject: {args.subject_code}, Exam: {args.exam_type}, Section: {args.section} ---", file=sys.stderr)
            output_result = {"status": "info", "message": "Class marksheet CLI generation not yet fully implemented."}
            # Potentially, you'd loop through all students matching class criteria and call generate_student_result_service
            # or a modified version, then aggregate results.

        print(json.dumps(output_result, indent=2, default=str)) # Ensure datetime is converted to str for JSON
    except ValueError as ve:
        print(json.dumps({"status": "error_cli_value", "message": str(ve)}, indent=2), file=sys.stdout)
        print(f"CLI Value Error: {ve}", file=sys.stderr)
        # traceback.print_exc(file=sys.stderr) # Uncomment for full traceback if needed
    except Exception as e:
        print(json.dumps({"status": "error_cli_exception", "message": str(e)}, indent=2), file=sys.stdout)
        print(f"CLI Unhandled Error: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr) # Always print traceback for unexpected errors
    finally:
        if 'client' in globals() and client: # Check if client was defined
            try:
                client.close()
                print("INFO (MarksheetGen CLI): MongoDB connection closed.", file=sys.stderr)
            except Exception as e_close:
                print(f"WARNING (MarksheetGen CLI): Error closing MongoDB connection: {e_close}", file=sys.stderr)